import os
import re
import requests
import tempfile
import time
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from deep_translator import GoogleTranslator
from tqdm import tqdm

# Initialize Google Translator
translator = GoogleTranslator(source='ja', target='en')

def needs_translation(text):
    """Check if text contains Japanese characters."""
    return bool(re.search(r'[\u3040-\u30ff\u3400-\u4dbf\u4e00-\u9fff]', text))

def translate_safe(text, retries=3):
    """Translate text with retry logic and delay to avoid rate limiting."""
    if not text.strip() or not needs_translation(text):
        return text
    
    for attempt in range(retries):
        try:
            # Sleep to prevent hitting Google Translate rate limits
            time.sleep(0.3)
            translated = translator.translate(text)
            return translated if translated else text
        except Exception as e:
            print(f"\n[Warning] Translation error on attempt {attempt+1} for text '{text[:20]}...': {e}")
            time.sleep(2) # Wait longer before retrying
            
    return text # Return original text if all retries fail

def process_table(document, table_html):
    """Parse HTML table and create a Word table."""
    soup = BeautifulSoup(table_html, 'html.parser')
    trs = soup.find_all('tr')
    
    if not trs:
        return
        
    grid = []
    for tr in trs:
        tds = tr.find_all(['td', 'th'])
        grid.append([td.get_text(strip=True) for td in tds])
        
    max_cols = max([len(row) for row in grid]) if grid else 1
    table = document.add_table(rows=len(grid), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(grid):
        for j, text in enumerate(row):
            if text:
                translated = translate_safe(text)
                try:
                    table.cell(i, j).text = translated
                except Exception as e:
                    pass

def process_markdown_to_docx(input_file, output_file):
    document = Document()
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_html = ""
    temp_dir = tempfile.mkdtemp()
    
    for line in tqdm(lines, desc="Processing Document"):
        line_stripped = line.strip()
        
        if not line_stripped:
            continue
            
        # --- Handle Table State ---
        if '<table' in line_stripped and not in_table:
            in_table = True
            table_html = line_stripped
            if '</table>' in line_stripped:
                process_table(document, table_html)
                in_table = False
                table_html = ""
            continue
            
        if in_table:
            table_html += " " + line_stripped
            if '</table>' in line_stripped:
                process_table(document, table_html)
                in_table = False
                table_html = ""
            continue
            
        # --- Handle Images ---
        if '<img' in line_stripped and 'src=' in line_stripped:
            soup = BeautifulSoup(line_stripped, 'html.parser')
            img_tag = soup.find('img')
            if img_tag and img_tag.get('src'):
                img_url = img_tag['src']
                try:
                    # Download image
                    response = requests.get(img_url, stream=True)
                    if response.status_code == 200:
                        img_path = os.path.join(temp_dir, f"temp_img_{len(document.paragraphs)}.jpg")
                        with open(img_path, 'wb') as img_f:
                            for chunk in response.iter_content(1024):
                                img_f.write(chunk)
                        
                        # Add image to doc
                        p = document.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run()
                        try:
                            # Parse width if exists to size it properly
                            width_pct = img_tag.get('width', '')
                            if '%' in width_pct:
                                pct = float(width_pct.replace('%', ''))
                                # Assume max width is 6.0 inches
                                r.add_picture(img_path, width=Inches(6.0 * (pct / 100.0)))
                            else:
                                r.add_picture(img_path, width=Inches(5.0))
                        except Exception:
                            r.add_picture(img_path, width=Inches(5.0))
                except Exception as e:
                    print(f"\nFailed to load image {img_url[:30]}...: {e}")
            continue

        # --- Handle Headings ---
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line_stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            translated_text = translate_safe(text)
            document.add_heading(translated_text, level=level)
            continue
            
        # --- Handle HTML text (like centered captions) ---
        if line_stripped.startswith('<div') and '</div>' in line_stripped:
            soup = BeautifulSoup(line_stripped, 'html.parser')
            text = soup.get_text(strip=True)
            if text:
                translated_text = translate_safe(text)
                p = document.add_paragraph(translated_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # --- Handle regular text ---
        translated_text = translate_safe(line_stripped)
        document.add_paragraph(translated_text)
        
    document.save(output_file)
    print(f"\nSuccess! Document saved to {output_file}")

if __name__ == "__main__":
    input_md = "jpy book 280 to 483 pg.pdf_by_PaddleOCR-VL-1.5.md"
    output_docx = "Translated_English_Output_pg280-483.docx"

    print(f"Starting conversion of {input_md}...")
    process_markdown_to_docx(input_md, output_docx)
