import os
import re
import requests
import tempfile
import time
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from deep_translator import GoogleTranslator
from langdetect import detect, DetectorFactory

# Ensure consistent results for language detection
DetectorFactory.seed = 0

# Initialize Google Translator
translator = GoogleTranslator(source='ja', target='en')

def detect_language(text):
    """Detect language of the text."""
    try:
        # Filter out numbers and special characters to improve detection
        clean_text = re.sub(r'[0-9\W_]+', '', text)
        if len(clean_text) < 5:
            return "unknown"
        return detect(clean_text)
    except:
        return "unknown"

def needs_translation(text):
    """Check if text contains Japanese characters."""
    return bool(re.search(r'[\u3040-\u30ff\u3400-\u4dbf\u4e00-\u9fff]', text))

def translate_safe_generator(text, retries=3):
    """Translate text with retry logic and yield logs on failure."""
    if not text.strip() or not needs_translation(text):
        return text, None
    
    for attempt in range(retries):
        try:
            time.sleep(0.3)
            translated = translator.translate(text)
            return translated if translated else text, None
        except Exception as e:
            log_msg = f"Translation error on attempt {attempt+1} for text '{text[:20]}...': {e}"
            time.sleep(2)
            if attempt < retries - 1:
                return None, log_msg
                
    return text, f"Failed translation for '{text[:20]}...' after {retries} attempts"

def process_table_stream(document, table_html):
    """Parse HTML table and create a Word table with translated content."""
    soup = BeautifulSoup(table_html, 'html.parser')
    trs = soup.find_all('tr')
    
    if not trs:
        return None
        
    grid = []
    for tr in trs:
        tds = tr.find_all(['td', 'th'])
        grid.append([td.get_text(strip=True) for td in tds])
        
    max_cols = max([len(row) for row in grid]) if grid else 1
    table = document.add_table(rows=len(grid), cols=max_cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(grid):
        for j, text in enumerate(row):
            if text:
                translated_text = None
                while translated_text is None:
                    res, log = translate_safe_generator(text)
                    if log: yield {"log": log}
                    if res is not None: translated_text = res
                
                try:
                    table.cell(i, j).text = translated_text
                except:
                    pass
    yield {"log": f"Processed table with {len(grid)} rows."}

def process_markdown_to_docx_stream(input_file, output_file, cancel_check=None):
    document = Document()
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Detect language from first few non-empty lines
    sample_text = ""
    for line in lines:
        if line.strip():
            sample_text += line.strip() + " "
            if len(sample_text) > 200: break
    
    detected_lang = detect_language(sample_text)
    yield {"stage": "Language Detection", "progress": 5, "log": f"Detected language: {detected_lang}", "detected_lang": detected_lang}

    in_table = False
    table_html = ""
    total_lines = len(lines)
    
    for idx, line in enumerate(lines):
        # Check for cancellation
        if cancel_check and cancel_check():
            yield {"stage": "Cancelled", "log": "Processing cancelled by user."}
            return

        line_stripped = line.strip()
        current_progress = 10 + int((idx / total_lines) * 85)
        
        if idx % 10 == 0:
            yield {"stage": "Translating Document", "progress": current_progress, "log": f"Processing line {idx}/{total_lines}..."}
            
        if not line_stripped and not in_table:
            continue
            
        # --- Handle Table State ---
        if '<table' in line_stripped and not in_table:
            in_table = True
            table_html = line_stripped
            if '</table>' in line_stripped:
                for update in process_table_stream(document, table_html):
                    yield update
                in_table = False
                table_html = ""
            continue
            
        if in_table:
            table_html += " " + line_stripped
            if '</table>' in line_stripped:
                for update in process_table_stream(document, table_html):
                    yield update
                in_table = False
                table_html = ""
            continue
            
        # --- Handle Headings ---
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line_stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            
            translated_text = None
            while translated_text is None:
                res, log = translate_safe_generator(text)
                if log: yield {"log": log}
                if res is not None: translated_text = res
                
            document.add_heading(translated_text, level=level)
            continue
            
        # --- Handle regular text ---
        translated_text = None
        while translated_text is None:
            res, log = translate_safe_generator(line_stripped)
            if log: yield {"log": log}
            if res is not None: translated_text = res
            
        document.add_paragraph(translated_text)
        
    yield {"stage": "Saving DOCX", "progress": 98, "log": "Formatting final Word Document..."}
    document.save(output_file)
    yield {"stage": "Completed", "progress": 100, "log": f"Success! Document saved to {output_file}"}
